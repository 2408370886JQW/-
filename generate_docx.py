import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(source_dir, output_filename, structure_map_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('产品交互原型评审文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('生成日期: 2026-02-13').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Product Structure Map
    doc.add_heading('1. 产品全链路结构图', level=1)
    doc.add_paragraph('下图展示了社交生活 App 的六大核心链路及页面层级关系：')
    
    if os.path.exists(structure_map_path):
        doc.add_picture(structure_map_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph('[结构图文件缺失]')
    
    doc.add_page_break()

    # Structure Definition
    structure = {
        "01_Meet_Flow": {
            "title": "2. 相见链路",
            "desc": "核心流程：用户发起邀约 -> 选择商家 -> 支付套餐 -> 完成订单。",
            "pages": [
                ("01_Relation_Selection", "2.1 相见页-01-关系选择", "用户选择邀约对象的关系类型（如朋友、恋人），系统据此推荐不同氛围的商家。"),
                ("02_Venue_List", "2.2 相见页-02-商家列表", "展示符合条件的商家列表，支持按距离、评分筛选。"),
                ("03_Venue_Detail", "2.3 相见页-03-商家详情", "展示商家环境图、特色菜品及用户评价。"),
                ("04_Package_Selection", "2.4 相见页-04-套餐选择", "用户选择具体的餐饮或活动套餐。"),
                ("05_Payment", "2.5 相见页-05-支付页", "确认订单金额，选择支付方式（微信/支付宝）。"),
                ("06_Payment_Success", "2.6 相见页-06-支付完成引导", "支付成功后的反馈页面，引导用户查看订单或发起导航。"),
                ("07_Order_Detail", "2.7 相见页-07-订单详情", "展示订单状态、核销码及商家信息。")
            ]
        },
        "02_Encounter_Flow": {
            "title": "3. 偶遇链路",
            "desc": "基于地理位置的社交发现功能。",
            "pages": [
                ("01_Map_Home", "3.1 偶遇页-01-地图首页", "地图模式展示附近的用户和活动热点。"),
                ("02_Filter_Modal", "3.2 偶遇页-02-筛选弹窗", "筛选地图上显示的用户性别、年龄段及活动类型。")
            ]
        },
        "03_Friends_Flow": {
            "title": "4. 好友链路",
            "desc": "管理已建立社交关系的好友列表。",
            "pages": [
                ("01_Friend_List", "4.1 好友页-01-列表模式", "展示好友列表，支持搜索和首字母索引。")
            ]
        },
        "04_Moments_Flow": {
            "title": "5. 动态链路",
            "desc": "用户分享生活点滴的内容社区。",
            "pages": [
                ("01_Moments_Feed", "5.1 动态页-01-动态流", "瀑布流展示好友及推荐的动态内容。"),
                ("02_Post_Moment", "5.2 动态页-02-发布动态", "图文发布编辑器，支持添加位置和话题。")
            ]
        },
        "05_Messages_Flow": {
            "title": "6. 消息链路",
            "desc": "即时通讯模块。",
            "pages": [
                ("01_Message_List", "6.1 消息页-01-消息列表", "展示所有私聊和群聊会话列表。")
            ]
        },
        "06_Profile_Flow": {
            "title": "7. 我的链路",
            "desc": "个人信息及设置管理。",
            "pages": [
                ("01_Profile_Center", "7.1 我的页-01-个人中心", "展示用户头像、资料编辑入口及功能设置。")
            ]
        }
    }

    for folder_key, content in structure.items():
        # Section Header
        doc.add_heading(content["title"], level=1)
        doc.add_paragraph(content["desc"])
        
        for page_folder, page_title, page_desc in content["pages"]:
            # Page Header
            doc.add_heading(page_title, level=2)
            doc.add_paragraph(page_desc)
            
            img_path = os.path.join(source_dir, folder_key, page_folder, "preview.png")
            
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(3.5)) # Adjust width to fit nicely
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f"[图片加载失败: {e}]")
            else:
                doc.add_paragraph(f"[图片未找到: {img_path}]")
            
            doc.add_paragraph('') # Add some spacing

    doc.save(output_filename)
    print(f"Word document generated: {output_filename}")

if __name__ == "__main__":
    source = "/home/ubuntu/social-life-app/product_package_en"
    output = "/home/ubuntu/social-life-app/产品交互原型评审文档_飞书版.docx"
    structure_map = "/home/ubuntu/social-life-app/product_structure.png"
    create_docx(source, output, structure_map)
